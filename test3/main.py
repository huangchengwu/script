import json
import requests
import time
from docx import Document
from jinja2 import Template
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import shutil
import os

# 获取当前时间戳
timestamp = int(time.time())

# 根据时间戳创建文件夹
folder_name = str(timestamp)
os.mkdir(folder_name)


# 生成文档
def new_docx(a, b, c, title):
    print("生成文档", title + ".docx", "进行中")

    source_file = "测试表格.docx"
    destination_file = os.path.join(folder_name+"/"+title + ".docx")
    # 复制文件
    shutil.copyfile(source_file, destination_file)
    # 打开现有文档
    doc = Document(folder_name+"/"+title  + ".docx")
    for paragraph in doc.paragraphs:
        template = Template(paragraph.text)
        output = template.render({"title": title})

        paragraph.text = output

        # 设置字体为宋体
        paragraph.runs[0].font.name = "宋体"

        # 设置字号为小二
        paragraph.runs[0].font.size = Pt(18)

        # 设置加粗样式
        paragraph.runs[0].bold = True

        # 设置对齐方式为居中
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table = doc.tables[0]

    template = Template(table.cell(0, 2).text)
    output = template.render({"title": title})
    # 活动名称
    cell=table.cell(0, 2)
    cell.text = output
    paragraph = cell.paragraphs[0]
    # 设置对齐方式为居中
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # 设置段落中所有文本的字体为宋体
    for run in paragraph.runs:
        run.font.name = '宋体'
        run.font.size = Pt(12)
    

    # 目标
    cell=table.cell(2, 1)
    cell.text = a
    paragraph = cell.paragraphs[0]
    # 设置对齐方式为居中
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # 设置段落中所有文本的字体为宋体
    for run in paragraph.runs:
        run.font.name = '宋体'
        run.font.size = Pt(12)
    
    # 收集

    cell=table.cell(3, 1)
    cell.text = b
    paragraph = cell.paragraphs[0]
    # 设置对齐方式为居中
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # 设置段落中所有文本的字体为宋体
    for run in paragraph.runs:
        run.font.name = '宋体'
        run.font.size = Pt(12)

    # 活动
    cell=table.cell(4, 1)
    cell.text = c
    paragraph = cell.paragraphs[0]
    # 设置对齐方式为居中
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # 设置段落中所有文本的字体为宋体
    for run in paragraph.runs:
        run.font.name = '宋体'
        run.font.size = Pt(12)


    print("生成文档", title + ".docx", "完成")

    doc.save(folder_name+"/"+title  + ".docx")


# 发消息
def do_task_send_text(token, payload):
    url = "https://t.hitosea.com/api/dialog/msg/sendtext"
    json_payload = json.dumps(payload)
    headers = {"Content-Type": "application/json", "Token": token}
    response = requests.post(url, data=json_payload, headers=headers)

    if response.status_code != 200:
        print("Error sending HTTP request:", response.text)
        return 0, response.status_code

    res_map = json.loads(response.text)
    dialog_id = res_map["data"]["id"] + 1
    return dialog_id, None


# 会话消息
def do_task_one(token, payload, ids, code, title):
    # print("会话查询",token, payload, ids, code, title)
    url = "https://t.hitosea.com/api/dialog/one"
    json_payload = json.dumps(payload)
    headers = {"Content-Type": "application/json", "Token": token}
    response = requests.post(url, data=json_payload, headers=headers)

    if response.status_code != 200:
        print("Error sending HTTP request:", response.text)
        return "", response.status_code

    res_map = json.loads(response.text)

    id = res_map["data"]["last_msg"]["id"]
    i = 0
    if code == 0:
        while True:
            msg = do_task_one(token, payload, ids, 1, title)
            if msg == "...":
                print("执行中", title, msg)
            else:
                m = json.loads(msg)
                new_docx(m["预期目标"], m["资源收集及利用"], m["活动流程"], title)

                break
            time.sleep(1)
            i += 1
    if id == ids:
        text = res_map["data"]["last_msg"]["msg"]["text"]
        return text
    return "", None


# 登陆
def do_task_login(payload):
    url = "https://t.hitosea.com/api/users/login"
    json_payload = json.dumps(payload)
    headers = {"Content-Type": "application/json"}
    response = requests.post(url, data=json_payload, headers=headers)

    if response.status_code != 200:
        print("Error sending HTTP request:", response.text)
        return "", response.status_code

    res_map = json.loads(response.text)
    res_to = f"{res_map['data']['token']}"
    return res_to, None


login_payload = {"email": "17710136904@163.com", "password": "Qq751164212."}
# login_payload = {"email": "aipaw@qq.com", "password": "Ab123456.."}


token, _ = do_task_login(login_payload)


text_content = """
2023年秋季学期大二班{{ title }}教育活动方案 请写一个教案格式如下
预期目标:
        内容
资源收集及利用:
        内容
活动流程: 
        内容

{{ describe }}  
然后给我返回格式为json 
预期目标 为key 
资源收集及利用 为key 
活动流程 为key 
json格式必须符合规范尽量是一行 并且内容详细要800字
"""

template = Template(text_content)
data = [
    {"title": "防空", "describe": ""},
    # {"title": "户外安全", "describe": "主要讲 防走丢 迷路了怎么办等"},
    # {"title": "危险的物品", "describe": "主要讲 小班异物入体，中班尖利的东西会伤人，大班暴力玩具"},
    # {"title": "防食物中毒", "describe": ""},
    # {"title": "防拐骗", "describe": ""},
]

dialog_ids = "9966"
for i in data:
    output = template.render(i)
    sendtext_payload = {"dialog_id": dialog_ids, "text": output}
    dialog_id, _ = do_task_send_text(token, sendtext_payload)

    one_payload = {
        "dialog_id": dialog_ids,
    }

    res = do_task_one(token, one_payload, dialog_id, 0, i["title"])
