from docx import Document
from jinja2 import Template
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import shutil
import os
source_file = "测试表格.docx"
destination_file = os.path.join("新测试表格.docx")
# 复制文件
shutil.copyfile(source_file, destination_file)
# 打开现有文档
doc = Document("新测试表格.docx")
for paragraph in doc.paragraphs:
    template = Template(paragraph.text)
    output = template.render({"title": "防溺水"})

    paragraph.text = output

    # 设置字体为宋体
    paragraph.runs[0].font.name = "宋体"

    # 设置字号为小二
    paragraph.runs[0].font.size = Pt(18)

    # 设置加粗样式
    paragraph.runs[0].bold = True

    # 设置对齐方式为居中
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
table = doc.tables[0]

template = Template(table.cell(0, 2).text)
output = template.render({"title": "防溺水"})
# 活动名称
cell=table.cell(0, 2)
cell.text = output
paragraph = cell.paragraphs[0]
# 设置对齐方式为居中
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
# 设置段落中所有文本的字体为宋体
for run in paragraph.runs:
    run.font.name = '宋体'
    run.font.size = Pt(12)
 

table = doc.tables[0]
# 目标
cell_text = table.cell(2, 1).text
# 收集
cell_text = table.cell(3, 1).text
# 活动
cell_text = table.cell(4, 1).text

print(cell_text)

doc.save("新测试表格.docx")
